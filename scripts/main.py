"""
Elevator Daily Digest - Main entry point.
Fetches elevator industry news, summarizes with AI, and saves outputs.
Supports daily mode (default) and weekly mode (--weekly or every Monday).
"""

import json
import os
import re
import sys
from datetime import datetime, timezone, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from sources import fetch_all
from summarize import summarize, summarize_weekly
from generate_site import generate_site
from audio import generate_audio, prune_old_audio


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Add a clickable hyperlink to a Word paragraph."""
    # python-docx doesn't expose hyperlinks natively, use lxml
    from docx.oxml.ns import qn
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = paragraph._p.makeelement(qn('w:hyperlink'), {qn('r:id'): r_id})
    run = hyperlink.makeelement(qn('w:r'), {})
    rPr = run.makeelement(qn('w:rPr'), {})
    rFonts = rPr.makeelement(qn('w:rFonts'), {qn('w:ascii'): 'Microsoft YaHei', qn('w:eastAsia'): 'Microsoft YaHei'})
    rPr.append(rFonts)
    sz = rPr.makeelement(qn('w:sz'), {qn('w:val'): '20'})  # 10pt in half-points
    rPr.append(sz)
    c = rPr.makeelement(qn('w:color'), {qn('w:val'): '0563C1'})
    rPr.append(c)
    u = rPr.makeelement(qn('w:u'), {qn('w:val'): 'single'})
    rPr.append(u)
    run.append(rPr)
    t = run.makeelement(qn('w:t'), {})
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_markdown_text(paragraph, text: str, font_size: Pt = Pt(10), color=None) -> None:
    """Add text with inline Markdown bold and links converted to Word formatting."""
    # Combined pattern: **bold** or [text](url)
    pattern = re.compile(r'\*\*(.+?)\*\*|\[([^\]]+)\]\(([^)]+)\)')
    last_end = 0
    for m in pattern.finditer(text):
        # Add plain text before the match
        prefix = text[last_end:m.start()]
        if prefix:
            run = paragraph.add_run(prefix)
            run.font.size = font_size
            if color:
                run.font.color.rgb = color
        if m.group(1) is not None:
            # Bold: **text**
            run = paragraph.add_run(m.group(1))
            run.bold = True
            run.font.size = font_size
            if color:
                run.font.color.rgb = color
        else:
            # Link: [text](url)
            _add_hyperlink(paragraph, m.group(2), m.group(3))
        last_end = m.end()
    # Add remaining text after the last match
    suffix = text[last_end:]
    if suffix:
        run = paragraph.add_run(suffix)
        run.font.size = font_size
        if color:
            run.font.color.rgb = color


def _generate_word(markdown: str, filepath: Path) -> None:
    """Convert the digest Markdown to a formatted Word document."""
    doc = Document()

    # Set default font (Latin + East-Asian)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Also set heading styles to Microsoft YaHei
    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Microsoft YaHei'
        hs.font.size = Pt(16 - i * 2)  # H1=14, H2=12, H3=10
        hs.font.color.rgb = RGBColor(0x1F, 0x23, 0x28)
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    lines = markdown.splitlines()
    title_text = ""
    i = 0

    # Extract title from first line
    if lines and lines[0].startswith('# '):
        title_text = lines[0][2:].strip()
        i = 1

    # Add document title
    if title_text:
        title_para = doc.add_heading(title_text, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add a blank line after title
        doc.add_paragraph()

    while i < len(lines):
        line = lines[i].strip()

        # Skip footer
        if line.startswith('---') or line.startswith('*Generated'):
            i += 1
            continue

        # Category heading: ## ⚠️ xxx
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue

        # Observation heading: ### xxx
        if line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            i += 1
            continue

        # Item title: - **title**：description or legacy - **[title]**：description
        if line.startswith('- **'):
            content = line[2:]  # remove "- " prefix
            m = re.match(r'\*\*\[?([^\]*]+?)\]?\*\*[：:]?\s*(.*)', content)
            if m:
                title_text = m.group(1).strip()
                subtitle = m.group(2).strip()
            else:
                title_text = content.strip()
                subtitle = ''
            p = doc.add_paragraph()
            p.style = doc.styles['List Bullet']
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(11)
            if subtitle:
                p.add_run(f" {subtitle}")
            i += 1
            continue

        # 来源/出典 line - format: 来源：URL（来源名称）
        if line.startswith('- 来源') or line.startswith('- 出典'):
            text = line[2:].strip()
            # Extract source name and URL from [name](url)
            m = re.match(r'(来源|出典)[：:]\s*\[([^\]]+)\]\(([^)]+)\)', text)
            if m:
                prefix = m.group(1)
                src_name = m.group(2)
                src_url = m.group(3)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                # 来源：URL（名称）
                run = p.add_run(f"{prefix}：")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                # URL as clickable hyperlink
                _add_hyperlink(p, src_url, src_url)
                # （来源名称）
                run = p.add_run(f"（{src_name}）")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                _add_markdown_text(p, text, Pt(10), RGBColor(0x66, 0x66, 0x66))
            i += 1
            continue

        # Other metadata lines (AI摘要/AI要約, 日期/日付)
        if any(line.startswith(p) for p in ['- AI摘要', '- 日期',
                                              '- AI要約', '- 日付']):
            text = line[2:].strip()
            # Strip "AI摘要：" / "AI要約：" prefix from display
            text = re.sub(r'^(?:AI摘要|AI要約)[：:]\s*', '', text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            _add_markdown_text(p, text, Pt(10), RGBColor(0x66, 0x66, 0x66))
            i += 1
            continue

        # Observation paragraph text
        if line and not line.startswith('#'):
            p = doc.add_paragraph()
            _add_markdown_text(p, line, Pt(11))
            i += 1
            continue

        i += 1

    doc.save(str(filepath))
    print(f"  Word document saved to {filepath}")


def main():
    beijing_tz = timezone(timedelta(hours=8))
    today = datetime.now(beijing_tz)
    date_str = today.strftime("%Y-%m-%d")

    # Always weekly mode
    week_end = today
    week_start = today - timedelta(days=7)
    week_label = f"{week_start.strftime('%Y-%m-%d')} ~ {week_end.strftime('%Y-%m-%d')}"
    print(f"=== Elevator Weekly Digest for {week_label} ===\n")

    # Step 1: Fetch data
    print("[Step 1] Fetching data from sources...")
    data = fetch_all()

    total = sum(len(v) for v in data.values())
    print(f"\nTotal items fetched: {total}")

    if total == 0:
        print("No items fetched. Exiting.")
        sys.exit(0)

    # Step 2: Save raw data
    data_dir = Path(__file__).parent.parent / "data"
    data_dir.mkdir(exist_ok=True)
    raw_file = data_dir / f"{date_str}.raw.json"
    raw_payload = {
        "date": date_str,
        "generated_at": today.isoformat(),
        "counts": {key: len(value) for key, value in data.items()},
        "items": data,
    }
    raw_file.write_text(
        json.dumps(raw_payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(f"\n[Step 2] Saved raw data to {raw_file}")

    # Filter out accident/safety-incident items — 不纳入日报
    data.pop('accident', None)

    # Step 3: AI summarization (weekly)
    print("\n[Step 3] Generating AI weekly summary (zh + ja)...")
    label_for_file = f"{week_start.strftime('%Y%m%d')}-{week_end.strftime('%Y%m%d')}"
    markdown, markdown_ja = summarize_weekly(data, week_label)

    # Step 4: Save Markdown (zh + ja)
    output_dir = Path(__file__).parent.parent / "daily"
    output_dir.mkdir(exist_ok=True)
    output_file = output_dir / f"{label_for_file}.md"
    output_file.write_text(markdown, encoding="utf-8")
    print(f"\n[Step 4] Saved to {output_file}")

    if markdown_ja:
        output_file_ja = output_dir / f"{label_for_file}.ja.md"
        output_file_ja.write_text(markdown_ja, encoding="utf-8")
        print(f"          Saved Japanese to {output_file_ja}")
    else:
        print("          No Japanese version generated.")
    print(f"\n[Step 4] Saved to {output_file}")

    # Step 5: Generate Word documents (zh + ja)
    print("\n[Step 5] Generating Word documents...")
    word_dir = Path(__file__).parent.parent / "weekly"
    word_dir.mkdir(exist_ok=True)
    word_file = word_dir / f"{label_for_file}.docx"
    _generate_word(markdown, word_file)
    if markdown_ja:
        word_file_ja = word_dir / f"{label_for_file}_ja.docx"
        _generate_word(markdown_ja, word_file_ja)

    # Step 6: Voice broadcast
    print("\n[Step 6] Generating voice broadcast...")
    audio_dir = Path(__file__).parent.parent / "docs" / "audio"
    try:
        generate_audio(markdown, label_for_file, audio_dir)
        prune_old_audio(audio_dir, keep=15)
    except Exception as e:
        print(f"  Audio generation skipped: {e}")

    # Step 7: Rebuild static site
    print("\n[Step 7] Rebuilding static site...")
    generate_site(root=Path(__file__).parent.parent)
    print("  Site rebuilt → docs/")

    print("Done!")


if __name__ == "__main__":
    main()
